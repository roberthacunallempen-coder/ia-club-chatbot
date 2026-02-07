from pathlib import Path
import PyPDF2
import docx
from typing import List
from sqlalchemy.orm import Session
from app.models.knowledge import Knowledge
from app.models.document import DocumentType
import logging
import re

logger = logging.getLogger(__name__)


async def parse_document(file_path: str, file_type: str) -> str:
    """
    Extract text from document
    
    Args:
        file_path: Path to document file
        file_type: Type of document (pdf, docx, txt, csv)
    
    Returns:
        Extracted text content
    """
    try:
        if file_type == "pdf":
            return await parse_pdf(file_path)
        elif file_type == "docx":
            return await parse_docx(file_path)
        elif file_type == "txt":
            return await parse_txt(file_path)
        elif file_type == "csv":
            return await parse_csv(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    except Exception as e:
        logger.error(f"Error parsing document {file_path}: {e}")
        raise


async def parse_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    text_content = []
    
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}")
            except Exception as e:
                logger.warning(f"Error extracting page {page_num + 1}: {e}")
                continue
    
    return "\n\n".join(text_content)


async def parse_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    doc = docx.Document(file_path)
    
    text_content = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_content.append(row_text)
    
    return "\n\n".join(text_content)


async def parse_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
        return file.read()


async def parse_csv(file_path: str) -> str:
    """Extract text from CSV file"""
    import csv
    
    text_content = []
    
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
        csv_reader = csv.reader(file)
        
        for row_num, row in enumerate(csv_reader):
            row_text = " | ".join(row)
            text_content.append(f"Row {row_num + 1}: {row_text}")
    
    return "\n".join(text_content)


async def create_knowledge_from_document(
    document_id: int,
    text: str,
    category: str,
    db: Session,
    chunk_size: int = 1000
) -> int:
    """
    Create knowledge items from document text
    Splits text into chunks and creates separate knowledge entries
    
    Args:
        document_id: Document ID
        text: Extracted text
        category: Category for knowledge items
        db: Database session
        chunk_size: Characters per chunk
    
    Returns:
        Number of knowledge items created
    """
    try:
        # Split text into chunks by paragraphs or sentences
        chunks = split_text_intelligently(text, chunk_size)
        
        knowledge_count = 0
        
        for idx, chunk in enumerate(chunks):
            if len(chunk.strip()) < 50:  # Skip very short chunks
                continue
            
            # Generate title from first line or sentence
            title = generate_title_from_text(chunk)
            
            # Extract keywords
            keywords = extract_keywords(chunk)
            
            # Create knowledge item
            knowledge = Knowledge(
                title=title,
                category=category,
                content=chunk.strip(),
                keywords=keywords,
                source=f"Document ID: {document_id} (Chunk {idx + 1})",
                priority=0,
                is_active=True
            )
            
            db.add(knowledge)
            knowledge_count += 1
        
        db.commit()
        logger.info(f"Created {knowledge_count} knowledge items from document {document_id}")
        
        return knowledge_count
    
    except Exception as e:
        logger.error(f"Error creating knowledge from document: {e}")
        db.rollback()
        raise


def split_text_intelligently(text: str, chunk_size: int = 1000) -> List[str]:
    """
    Split text into intelligent chunks
    Tries to split by paragraphs, then sentences
    """
    chunks = []
    
    # Split by double newlines (paragraphs)
    paragraphs = text.split('\n\n')
    
    current_chunk = ""
    
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        
        # If adding this paragraph exceeds chunk size
        if len(current_chunk) + len(paragraph) > chunk_size:
            if current_chunk:
                chunks.append(current_chunk)
                current_chunk = ""
            
            # If paragraph itself is too long, split by sentences
            if len(paragraph) > chunk_size:
                sentences = re.split(r'[.!?]+\s+', paragraph)
                for sentence in sentences:
                    if len(current_chunk) + len(sentence) > chunk_size:
                        if current_chunk:
                            chunks.append(current_chunk)
                        current_chunk = sentence
                    else:
                        current_chunk += " " + sentence if current_chunk else sentence
            else:
                current_chunk = paragraph
        else:
            current_chunk += "\n\n" + paragraph if current_chunk else paragraph
    
    if current_chunk:
        chunks.append(current_chunk)
    
    return chunks


def generate_title_from_text(text: str, max_length: int = 100) -> str:
    """Generate a title from text (first sentence or line)"""
    lines = text.split('\n')
    first_line = lines[0].strip()
    
    if first_line:
        # Take first sentence
        sentences = re.split(r'[.!?]+', first_line)
        title = sentences[0].strip()
        
        if len(title) > max_length:
            title = title[:max_length] + "..."
        
        return title if title else "Untitled Knowledge"
    
    return "Untitled Knowledge"


def extract_keywords(text: str, max_keywords: int = 10) -> List[str]:
    """
    Extract keywords from text (simple approach)
    For better results, could use NLP libraries like spaCy
    """
    # Remove common words
    stop_words = {
        'el', 'la', 'de', 'que', 'y', 'a', 'en', 'un', 'ser', 'se',
        'no', 'haber', 'por', 'con', 'su', 'para', 'como', 'estar',
        'tener', 'le', 'lo', 'todo', 'pero', 'más', 'hacer', 'o',
        'poder', 'decir', 'este', 'ir', 'otro', 'ese', 'si', 'me',
        'ya', 'ver', 'porque', 'dar', 'cuando', 'él', 'muy', 'sin',
        'the', 'be', 'to', 'of', 'and', 'a', 'in', 'that', 'have',
        'it', 'for', 'not', 'on', 'with', 'he', 'as', 'you', 'do'
    }
    
    # Extract words (alphanumeric, 3+ chars)
    words = re.findall(r'\b[a-záéíóúñA-ZÁÉÍÓÚÑ]{3,}\b', text.lower())
    
    # Count frequency
    word_freq = {}
    for word in words:
        if word not in stop_words:
            word_freq[word] = word_freq.get(word, 0) + 1
    
    # Get top keywords
    sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
    keywords = [word for word, freq in sorted_words[:max_keywords]]
    
    return keywords
